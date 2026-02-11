
import os
import sys
from datetime import date
# PyQt5 imports moved to inside methods to avoid dependency in web app
from db_adapter import connect_db, PSYCOPG2_AVAILABLE
from logger import log_error, log_info

class ReportsHandler:
    @staticmethod
    def _norm(t):
        if not t: return ""
        return t.replace("أ", "ا").replace("إ", "ا").replace("آ", "ا").replace("ة", "ه").replace("ى", "ي")

    @staticmethod
    def generate_word_for_dept(dept_name, save_path):
        """Generates a Word document for a specific department."""
        try:
            import docx
            from docx.shared import Pt, Inches, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            raise ImportError("مكتبة python-docx غير مثبتة. يرجى تثبيتها باستخدام pip install python-docx")

        try:
            # 1. Fetch Data
            conn = connect_db()
            c = conn.cursor()
            
            # Fetch guarantees for this department
            query = """
                SELECT bank, g_type, g_no, amount, beneficiary, project_name, end_date, user_status, cash_flag 
                FROM guarantees 
                WHERE department = ?
            """
            
            # Adjust query for Postgres if needed
            if PSYCOPG2_AVAILABLE:
                try:
                    import psycopg2.extensions
                    if isinstance(conn, psycopg2.extensions.connection):
                        query = query.replace('?', '%s')
                except ImportError:
                    pass
            
            c.execute(query, (dept_name,))
            rows = c.fetchall()
            conn.close()

            final_keywords = ["افراج", "إفراج", "مردود", "ملغى", "ملغي", "إلغاء", "الغاء", "مصادر", "مصادرة", "انتهاء الغرض", "منتهي"]
            
            data_by_bank = {} # bank -> list of rows
            
            for r in rows:
                bank, g_type, g_no, amount, beneficiary, project_name, end_date, u_status, cash_flag = r
                
                # Check Final
                is_final = False
                if u_status:
                    u_norm = ReportsHandler._norm(str(u_status))
                    for k in final_keywords:
                        if ReportsHandler._norm(k) in u_norm:
                            is_final = True
                            break
                
                if is_final:
                    continue
                
                # Treat as Active (including Unregistered)
                b_key = bank if bank else "غير محدد"
                if b_key not in data_by_bank:
                    data_by_bank[b_key] = []
                data_by_bank[b_key].append(r)

            if not data_by_bank:
                return False # No data to print

            # 2. Create Document
            doc = docx.Document()
            
            # --- Watermark Logic (Added) ---
            try:
                # Assuming logo is at static/images/logo.png
                base_path = getattr(sys, '_MEIPASS', os.path.abspath("."))
                logo_path = os.path.join(base_path, "static", "images", "logo.png")
                
                if os.path.exists(logo_path):
                    # Add to first section header
                    section = doc.sections[0]
                    header = section.header
                    header.is_linked_to_previous = False
                    
                    from docx.oxml import parse_xml
                    from docx.oxml.ns import nsdecls
                    
                    # Add image to header relationship
                    rId, _ = header.part.get_or_add_image(logo_path)
                    
                    # Create VML XML for Watermark
                    # Opacity="0.1" -> 10% Opacity (as requested: transparency 10% interpreted as faint watermark)
                    xml = f"""
                    <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                         xmlns:v="urn:schemas-microsoft-com:vml"
                         xmlns:o="urn:schemas-microsoft-com:office:office">
                      <w:pPr>
                        <w:pStyle w:val="Header"/>
                      </w:pPr>
                      <w:r>
                        <w:pict>
                          <v:shape id="WatermarkLogo" type="#_x0000_t75" 
                                   style="position:absolute;margin-left:0;margin-top:0;width:300pt;height:300pt;z-index:-251659264;mso-wrap-edited:f;mso-position-horizontal:center;mso-position-horizontal-relative:margin;mso-position-vertical:center;mso-position-vertical-relative:margin">
                            <v:imagedata r:id="{rId}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" o:title="Logo"/>
                            <v:fill on="t" opacity="10%"/>
                          </v:shape>
                        </w:pict>
                      </w:r>
                    </w:p>
                    """
                    header._element.append(parse_xml(xml))
            except Exception as e:
                log_error(f"Watermark addition failed: {e}")
            # -------------------------------

            # Disable spell check (Hide red lines)
            settings = doc.settings.element
            hideSpellingErrors = OxmlElement('w:hideSpellingErrors')
            hideSpellingErrors.set(qn('w:val'), 'true')
            settings.append(hideSpellingErrors)

            hideGrammarErrors = OxmlElement('w:hideGrammarErrors')
            hideGrammarErrors.set(qn('w:val'), 'true')
            settings.append(hideGrammarErrors)

            # --- Section Setup (Margins & Borders) ---
            section = doc.sections[0]
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

            # Page Borders
            sectPr = section._sectPr
            pgBorders = OxmlElement('w:pgBorders')
            pgBorders.set(qn('w:offsetFrom'), 'page')
            for border_name in ('top', 'left', 'bottom', 'right'):
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '12') # 1.5pt
                border.set(qn('w:space'), '24')
                border.set(qn('w:color'), 'auto')
                pgBorders.append(border)
            sectPr.append(pgBorders)
            
            # --- Footer ---
            footer = section.footer
            p_footer = footer.paragraphs[0]
            p_footer.text = "ادارة الضمانات - شركات ابو سرهد"
            p_footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_footer.style = doc.styles['Normal']
            p_footer.runs[0].font.size = Pt(10)
            p_footer.runs[0].font.bold = True
            
            # --- Header ---
            header = section.header
            p_header = header.paragraphs[0]
            p_header.text = f"ضمانات قسم {dept_name}"
            p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_header.style = doc.styles['Normal']
            run_header = p_header.runs[0]
            run_header.font.size = Pt(16)
            run_header.font.bold = True
            run_header.font.name = 'Arial'

            # Helper for RTL
            def set_rtl(paragraph):
                pPr = paragraph._p.get_or_add_pPr()
                pPr.set(qn('w:bidi'), '1')
                run = paragraph.runs[0] if paragraph.runs else None
                if run:
                    rPr = run._r.get_or_add_rPr()
                    rPr.set(qn('w:rtl'), '1')

            # Default Style
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(12)
            
            # Enable RTL for Normal style (Text direction from right to left)
            pPr = style.element.get_or_add_pPr()
            pPr.set(qn('w:bidi'), '1')
            rPr = style.element.get_or_add_rPr()
            rPr.set(qn('w:rtl'), '1')
            
            # Date (Top Right)
            p_date = doc.add_paragraph(date.today().strftime("%Y/%m/%d"))
            p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Address
            p_to = doc.add_paragraph(f"السادة/ {dept_name} المحترمين،")
            p_to.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_rtl(p_to)
            p_to.runs[0].bold = True
            p_to.runs[0].font.size = Pt(18)

            p_greet = doc.add_paragraph("تحية طيبة وبعد،")
            p_greet.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_rtl(p_greet)
            
            # Body
            long_text = "نود إفادتكم بإعداد كشف بالضمانات القائمة والمسجلة باسم القسم الخاص بكم، ونرجو منكم التكرم بمراجعة هذه الضمانات والتأكد من مدى استخدامها أو الحاجة الفعلية لها وإفادتنا بما اذا كان ما زال الضمان مستخدمًا حاليًا وفي حال عدم استخدامه هل يمكن إلغاؤه أو استرداده، ونأمل منكم تزويدنا بالرد في أقرب وقت ممكن لتحديث بيانات الضمانات الخاصة بكم"
            
            p_body = doc.add_paragraph(long_text)
            p_body.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_rtl(p_body)
            if p_body.runs:
                p_body.runs[0].bold = True
                p_body.runs[0].font.size = Pt(11)

            p_thanks = doc.add_paragraph("شاكرين تعاونكم الدائم")
            p_thanks.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_rtl(p_thanks)
            if p_thanks.runs:
                p_thanks.runs[0].bold = True
                p_thanks.runs[0].font.size = Pt(11)
            
            doc.add_paragraph("") # Spacer

            # Tables per Bank
            headers = ["نوع الضمان", "رقم الضمان", "مبلغ الضمان", "اسم الجهة", "اسم المشروع", "تاريخ الانتهاء"]
            
            # Calculate column widths
            col_widths = [
                Inches(0.7),  # Type
                Inches(1.1),  # G No
                Inches(1.1),  # Amount
                Inches(1.95), # Beneficiary
                Inches(1.95), # Project
                Inches(0.8)   # End Date
            ]

            grand_total = 0.0

            for bank, items in data_by_bank.items():
                # Bank Header
                # Use RLM (\u200f) to force RTL direction for the bullet and punctuation
                p_bank = doc.add_paragraph(f"\u200f•   الضمانات القائمة علي {bank}")
                p_bank.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                set_rtl(p_bank)
                p_bank.runs[0].bold = True
                p_bank.runs[0].font.size = Pt(14)

                # Table
                table = doc.add_table(rows=1, cols=6)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Enable RTL for Table (Columns direction from right to left)
                tblPr = table._tbl.tblPr
                bidiVisual = OxmlElement('w:bidiVisual')
                tblPr.append(bidiVisual)

                table.autofit = False 
                table.allow_autofit = False
                
                # Apply widths to columns
                for i, width in enumerate(col_widths):
                    table.columns[i].width = width

                # Fill Headers
                hdr_cells = table.rows[0].cells
                for i, h in enumerate(headers):
                    hdr_cells[i].text = h
                    hdr_cells[i].width = col_widths[i] # Explicitly set cell width
                    hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    hdr_cells[i].paragraphs[0].runs[0].bold = True
                    hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
                
                total_bank = 0.0
                
                for item in items:
                    # bank, g_type, g_no, amount, beneficiary, project_name, end_date
                    g_type = item[1] or ""
                    g_no = item[2] or ""
                    amount = item[3] or 0.0
                    bene = item[4] or ""
                    proj = item[5] or ""
                    end_d = item[6] or ""
                    
                    total_bank += amount
                    
                    row_cells = table.add_row().cells
                    vals = [g_type, g_no, f"{amount:,.2f}", bene, proj, end_d]
                    
                    for i, v in enumerate(vals):
                        row_cells[i].text = str(v)
                        row_cells[i].width = col_widths[i] # Explicitly set cell width
                        row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        row_cells[i].paragraphs[0].runs[0].bold = True
                        row_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
                        
                        # Ensure RTL for Arabic content in Project/Beneficiary
                        if i in [3, 4]:
                            set_rtl(row_cells[i].paragraphs[0])
                
                # Total Row
                row_total = table.add_row().cells
                row_total[1].text = "الإجمالي"
                row_total[2].text = f"{total_bank:,.2f}"
                
                # Set widths for total row
                for i in range(6):
                    row_total[i].width = col_widths[i]
                
                row_total[1].paragraphs[0].runs[0].bold = True
                row_total[2].paragraphs[0].runs[0].bold = True
                row_total[1].paragraphs[0].runs[0].font.size = Pt(9)
                row_total[2].paragraphs[0].runs[0].font.size = Pt(9)
                row_total[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_total[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                grand_total += total_bank
                doc.add_paragraph("")

            # Grand Total
            p_grand = doc.add_paragraph(f"الإجمالي الكلي للقسم: {grand_total:,.2f}")
            p_grand.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_rtl(p_grand)
            p_grand.runs[0].bold = True
            p_grand.runs[0].font.size = Pt(12)

            doc.save(save_path)
            log_info(f"Report generated successfully: {save_path}")
            return True

        except Exception as e:
            log_error(f"Error generating word report for {dept_name}: {e}")
            raise e

    @staticmethod
    def generate_dept_letter_ui(parent, dept_name):
        try:
            from PyQt5 import QtWidgets, QtCore
            # 1. Select File Location
            today_str = date.today().strftime("%Y-%m-%d")
            clean_dept = str(dept_name).replace("/", "-").replace("\\", "-")
            fname, _ = QtWidgets.QFileDialog.getSaveFileName(
                parent, "حفظ خطاب القسم", f"خطاب_{clean_dept}_{today_str}.docx", "Word Documents (*.docx)"
            )
            if not fname:
                return

            # 2. Generate
            try:
                success = ReportsHandler.generate_word_for_dept(dept_name, fname)
                
                if success:
                    QtWidgets.QMessageBox.information(parent, "نجاح", f"تم حفظ الخطاب بنجاح في:\n{fname}")
                    try:
                        os.startfile(fname)
                    except:
                        pass
                else:
                    QtWidgets.QMessageBox.information(parent, "تنبيه", "لا توجد ضمانات سارية لهذا القسم.")
            except ImportError:
                 QtWidgets.QMessageBox.critical(parent, "خطأ", "مكتبة python-docx غير مثبتة.\nيرجى تثبيتها باستخدام:\npip install python-docx")
            except Exception as e:
                QtWidgets.QMessageBox.critical(parent, "خطأ", f"حدث خطأ أثناء إنشاء الملف: {e}")

        except Exception as e:
            log_error(f"UI Error in generate_dept_letter: {e}")
            QtWidgets.QMessageBox.critical(parent, "خطأ", f"حدث خطأ: {e}")

    @staticmethod
    def generate_all_dept_letters_ui(parent):
        try:
            from PyQt5 import QtWidgets, QtCore
            # 1. Select Directory
            dir_path = QtWidgets.QFileDialog.getExistingDirectory(parent, "اختر مجلد لحفظ التقارير")
            if not dir_path:
                return

            # 2. Get All Departments
            conn = connect_db()
            c = conn.cursor()
            # Get distinct departments that have active guarantees
            c.execute("SELECT DISTINCT department FROM guarantees WHERE department IS NOT NULL AND department != ''")
            depts = [r[0] for r in c.fetchall()]
            conn.close()

            if not depts:
                QtWidgets.QMessageBox.information(parent, "تنبيه", "لا توجد أقسام مسجلة.")
                return

            # 3. Process
            progress = QtWidgets.QProgressDialog("جاري إنشاء التقارير...", "إلغاء", 0, len(depts), parent)
            progress.setWindowModality(QtCore.Qt.WindowModal)
            
            count = 0
            today_str = date.today().strftime("%Y-%m-%d")
            
            for i, dept in enumerate(depts):
                if progress.wasCanceled():
                    break
                
                progress.setValue(i)
                progress.setLabelText(f"جاري معالجة: {dept}")
                
                clean_dept = str(dept).replace("/", "-").replace("\\", "-")
                fname = os.path.join(dir_path, f"خطاب_{clean_dept}_{today_str}.docx")
                
                try:
                    if ReportsHandler.generate_word_for_dept(dept, fname):
                        count += 1
                except Exception as e:
                    log_error(f"Failed to generate report for {dept}: {e}")
            
            progress.setValue(len(depts))
            
            QtWidgets.QMessageBox.information(parent, "تم", f"تم إنشاء {count} ملف بنجاح في:\n{dir_path}")
            try:
                os.startfile(dir_path)
            except:
                pass

        except Exception as e:
            log_error(f"UI Error in generate_all_dept_letters: {e}")
            QtWidgets.QMessageBox.critical(parent, "خطأ", f"حدث خطأ: {e}")
